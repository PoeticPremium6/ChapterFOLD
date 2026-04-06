#!/usr/bin/env python3
from __future__ import annotations

import html
import re
import warnings
from dataclasses import dataclass
from pathlib import Path
from typing import List, Tuple

from bs4 import BeautifulSoup, XMLParsedAsHTMLWarning
from ebooklib import ITEM_DOCUMENT, epub
from docx import Document

from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

warnings.filterwarnings("ignore", category=XMLParsedAsHTMLWarning)

try:
    from weasyprint import HTML
except Exception as e:
    HTML = None
    WEASYPRINT_IMPORT_ERROR = e
else:
    WEASYPRINT_IMPORT_ERROR = None

DEFAULT_FONT_STACK = '"Garamond", "EB Garamond", "Cormorant Garamond", serif'


@dataclass
class LayoutSettings:
    trim_width_cm: float = 15.24
    trim_height_cm: float = 22.86
    margin_top_cm: float = 1.5
    margin_bottom_cm: float = 1.5
    margin_inside_cm: float = 1.8
    margin_outside_cm: float = 1.0
    font_size_pt: float = 11.5
    line_height: float = 1.35
    font_family: str = DEFAULT_FONT_STACK
    drop_notes: bool = False
    paragraph_spacing_mode: str = "traditional"


@dataclass
class CleanupSettings:
    join_soft_wrapped_lines: bool = True
    join_dialogue_continuations: bool = True
    merge_dialogue_paragraphs: bool = False
    aggressive_mode: bool = False
    collapse_extra_blank_lines: bool = True
    preserve_scene_breaks: bool = True


@dataclass
class EpubContent:
    detected_title: str
    detected_author: str
    sections: List[Tuple[str, str]]


@dataclass
class EpubToPdfResult:
    input_epub: Path
    output_dir: Path
    output_pdf: Path
    output_docx: Path | None
    output_markdown: Path | None
    book_slug: str
    detected_title: str
    detected_author: str
    used_title: str
    used_author: str


SCENE_BREAK_RE = re.compile(
    r"^\s*(\*\s*){3,}$|^\s*#{3,}\s*$|^\s*-\s*-\s*-\s*$|^\s*~\s*~\s*~\s*$"
)

DIALOGUE_TAG_START_RE = re.compile(
    r"^(?:"
    r"he|she|they|i|we|it|you|his|her|their|the|"
    r"said|asked|whispered|murmured|replied|answered|shouted|yelled|"
    r"cried|called|snapped|hissed|muttered|breathed|added|continued|"
    r"harry|ron|hermione|draco|ginny|luna|neville|snape|remus|sirius|"
    r"malfoy|potter|granger|he'd|she'd|they'd|he'll|she'll|they'll"
    r")\b",
    flags=re.IGNORECASE,
)

LOWER_CONTINUATION_RE = re.compile(r"^[a-z(\[\u2014\-']")
TERMINAL_END_RE = re.compile(r"""[.!?]["\u201d\u2019']?$""")
OPENING_QUOTE_RE = re.compile(r"""^["\u201c\u2018]""")


def cm(value: float) -> str:
    return f"{value:.3f}cm"


def slugify(value: str) -> str:
    value = value.strip().lower()
    value = re.sub(r"[^\w\s-]", "", value, flags=re.UNICODE)
    value = re.sub(r"[\s_-]+", "-", value)
    value = re.sub(r"-{2,}", "-", value)
    return value.strip("-") or "book"


def build_book_slug(
    *,
    title: str | None = None,
    author: str | None = None,
    fallback_stem: str | None = None,
) -> str:
    parts = []
    if author:
        parts.append(slugify(author))
    if title:
        parts.append(slugify(title))
    if parts:
        return "__".join(parts)
    if fallback_stem:
        return slugify(fallback_stem)
    return "book"


def output_dir_for_book(base_dir: Path, book_slug: str) -> Path:
    return base_dir / f"{book_slug}_output"


def interior_pdf_name(book_slug: str) -> str:
    return f"{book_slug}__interior.pdf"


def editable_docx_name(book_slug: str) -> str:
    return f"{book_slug}__editable.docx"


def markdown_name(book_slug: str) -> str:
    return f"{book_slug}__google-docs.md"


def normalize_line_endings(text: str) -> str:
    return text.replace("\r\n", "\n").replace("\r", "\n")


def normalize_spaces(text: str) -> str:
    text = text.replace("\xa0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r" *\n *", "\n", text)
    return text


def is_scene_break(line: str) -> bool:
    return bool(SCENE_BREAK_RE.match(line.strip()))


def join_dialogue_line_pair(current: str, nxt: str) -> str | None:
    current = current.rstrip()
    nxt = nxt.lstrip()
    if not current or not nxt:
        return None

    if re.search(r"""[,\u2014\-]|["\u201d\u2019]$""", current) and LOWER_CONTINUATION_RE.match(nxt):
        return f"{current} {nxt}"

    if re.search(r"""["\u201d\u2019]$""", current) and DIALOGUE_TAG_START_RE.match(nxt):
        return f"{current} {nxt}"

    return None


def clean_block_lines(lines: list[str], settings: CleanupSettings) -> list[str]:
    if not lines:
        return []

    result: list[str] = []
    i = 0
    while i < len(lines):
        current = lines[i].strip()
        if settings.join_dialogue_continuations and i + 1 < len(lines):
            joined = join_dialogue_line_pair(current, lines[i + 1])
            if joined is not None:
                result.append(joined)
                i += 2
                continue
        result.append(current)
        i += 1

    if settings.join_soft_wrapped_lines:
        merged = " ".join(x for x in result if x.strip())
        merged = re.sub(r" {2,}", " ", merged).strip()
        return [merged] if merged else []

    return result


def clean_text_block(text: str, settings: CleanupSettings | None = None) -> str:
    settings = settings or CleanupSettings()
    text = normalize_line_endings(text)
    text = normalize_spaces(text)
    lines = text.split("\n")

    output_blocks: list[str] = []
    current_block: list[str] = []

    def flush_current_block() -> None:
        nonlocal current_block
        if not current_block:
            return
        cleaned_lines = clean_block_lines(current_block, settings)
        if cleaned_lines:
            output_blocks.append("\n".join(cleaned_lines))
        current_block = []

    blank_run = 0
    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            blank_run += 1
            flush_current_block()
            if not settings.collapse_extra_blank_lines or blank_run == 1:
                output_blocks.append("")
            continue

        blank_run = 0
        if settings.preserve_scene_breaks and is_scene_break(line):
            flush_current_block()
            output_blocks.append(line)
            continue

        current_block.append(line)

    flush_current_block()

    cleaned = "\n".join(output_blocks)
    if settings.collapse_extra_blank_lines:
        cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def get_metadata(book: epub.EpubBook) -> Tuple[str, str]:
    title = "Unknown Title"
    author = "Unknown Author"

    title_meta = book.get_metadata("DC", "title")
    if title_meta and title_meta[0] and title_meta[0][0]:
        title = str(title_meta[0][0]).strip()

    creator_meta = book.get_metadata("DC", "creator")
    if creator_meta and creator_meta[0] and creator_meta[0][0]:
        author = str(creator_meta[0][0]).strip()

    return title, author


def load_epub_content(epub_path: Path) -> EpubContent:
    book = epub.read_epub(str(epub_path))
    title, author = get_metadata(book)
    sections: List[Tuple[str, str]] = []

    for item in book.get_items_of_type(ITEM_DOCUMENT):
        raw = item.get_content()
        soup = BeautifulSoup(raw, "xml")
        for bad in soup(["script", "style"]):
            bad.decompose()

        body = soup.find("body") or soup

        heading = ""
        for tag in body.find_all(["h1", "h2", "h3"]):
            txt = tag.get_text(" ", strip=True)
            if txt:
                heading = txt
                break

        body_html = "".join(str(x) for x in body.contents).strip()
        if body_html:
            sections.append((heading, body_html))

    return EpubContent(
        detected_title=title,
        detected_author=author,
        sections=sections,
    )


def _selectors_to_drop(drop_notes: bool) -> list[str]:
    selectors = [".landmark", ".actions", ".download", "nav"]
    if drop_notes:
        selectors.extend([
            ".preface",
            ".notes",
            ".endnotes",
            "#notes",
            "#work_endnotes",
        ])
    return selectors


def _drop_unwanted_nodes(soup: BeautifulSoup, *, drop_notes: bool) -> None:
    for selector in _selectors_to_drop(drop_notes):
        for el in soup.select(selector):
            el.decompose()


def _raw_paragraph_items_from_fragment(fragment: str, *, drop_notes: bool = False) -> list[tuple[str, str]]:
    soup = BeautifulSoup(fragment, "xml")
    _drop_unwanted_nodes(soup, drop_notes=drop_notes)
    body = soup.find("body") or soup

    items: list[tuple[str, str]] = []
    block_tags = {"p", "blockquote", "div", "li"}
    scene_tags = {"hr"}

    for node in body.descendants:
        if not getattr(node, "name", None):
            continue

        if node.name in scene_tags:
            items.append(("scene", "***"))
            continue

        if node.name not in block_tags:
            continue

        if node.find(block_tags):
            continue

        text = node.get_text("\n", strip=False)
        cleaned = clean_text_block(text)
        if cleaned.strip():
            for chunk in [part.strip() for part in cleaned.split("\n\n") if part.strip()]:
                if is_scene_break(chunk):
                    items.append(("scene", "***"))
                else:
                    items.append(("p", chunk))

    deduped: list[tuple[str, str]] = []
    previous = None
    for item in items:
        if item == previous and item[0] == "scene":
            continue
        deduped.append(item)
        previous = item
    return deduped


def should_merge_dialogue_paragraphs(
    current: str,
    nxt: str,
    settings: CleanupSettings,
) -> bool:
    current = current.strip()
    nxt = nxt.strip()
    if not current or not nxt:
        return False

    if is_scene_break(current) or is_scene_break(nxt):
        return False

    if re.search(r"""["\u201d\u2019]$""", current) and DIALOGUE_TAG_START_RE.match(nxt):
        return True

    if re.search(r"""[,;:\u2014\-]$|["\u201d\u2019]$""", current) and (
        LOWER_CONTINUATION_RE.match(nxt) or DIALOGUE_TAG_START_RE.match(nxt)
    ):
        return True

    if not settings.aggressive_mode:
        return False

    if len(current) <= 110 and not TERMINAL_END_RE.search(current):
        if LOWER_CONTINUATION_RE.match(nxt) or DIALOGUE_TAG_START_RE.match(nxt):
            return True

    if len(current) <= 90 and OPENING_QUOTE_RE.match(current):
        if not TERMINAL_END_RE.search(current):
            return True

    if len(nxt) <= 120 and (
        LOWER_CONTINUATION_RE.match(nxt) or DIALOGUE_TAG_START_RE.match(nxt)
    ):
        if not TERMINAL_END_RE.search(current):
            return True

    return False


def merge_adjacent_paragraph_items(
    items: list[tuple[str, str]],
    settings: CleanupSettings,
) -> list[tuple[str, str]]:
    if not settings.merge_dialogue_paragraphs:
        return items

    merged: list[tuple[str, str]] = []
    for kind, text in items:
        if (
            merged
            and kind == "p"
            and merged[-1][0] == "p"
            and should_merge_dialogue_paragraphs(merged[-1][1], text, settings)
        ):
            merged[-1] = ("p", f"{merged[-1][1].rstrip()} {text.lstrip()}")
        else:
            merged.append((kind, text))
    return merged


def extract_clean_items_from_html(
    fragment: str,
    *,
    drop_notes: bool = False,
    cleanup_settings: CleanupSettings | None = None,
) -> list[tuple[str, str]]:
    cleanup_settings = cleanup_settings or CleanupSettings()
    raw_items = _raw_paragraph_items_from_fragment(fragment, drop_notes=drop_notes)

    normalized: list[tuple[str, str]] = []
    for kind, text in raw_items:
        if kind == "scene":
            if cleanup_settings.preserve_scene_breaks:
                normalized.append(("scene", "***"))
            continue

        cleaned = clean_text_block(text, cleanup_settings)
        if cleaned.strip():
            normalized.append(("p", cleaned.strip()))

    return merge_adjacent_paragraph_items(normalized, cleanup_settings)


def sanitize_section_html(
    fragment: str,
    *,
    drop_notes: bool = False,
    cleanup_settings: CleanupSettings | None = None,
) -> str:
    items = extract_clean_items_from_html(
        fragment,
        drop_notes=drop_notes,
        cleanup_settings=cleanup_settings,
    )

    html_parts: list[str] = []
    for kind, text in items:
        if kind == "scene":
            html_parts.append('<hr class="scene-break" />')
        else:
            html_parts.append(f"<p>{html.escape(text)}</p>")
    return "\n".join(html_parts)


def extract_clean_text_from_html(
    fragment: str,
    *,
    drop_notes: bool = False,
    cleanup_settings: CleanupSettings | None = None,
) -> str:
    items = extract_clean_items_from_html(
        fragment,
        drop_notes=drop_notes,
        cleanup_settings=cleanup_settings,
    )

    parts: list[str] = []
    for kind, text in items:
        if kind == "scene":
            parts.append("***")
        else:
            parts.append(text)
    return "\n\n".join(parts).strip()


def build_clean_text_sections(
    sections: List[Tuple[str, str]],
    *,
    drop_notes: bool,
    cleanup_settings: CleanupSettings | None = None,
) -> List[Tuple[str, str]]:
    cleaned_sections: List[Tuple[str, str]] = []
    for heading, raw_html in sections:
        cleaned_text = extract_clean_text_from_html(
            raw_html,
            drop_notes=drop_notes,
            cleanup_settings=cleanup_settings,
        )
        if cleaned_text.strip():
            cleaned_sections.append((heading, cleaned_text))
    return cleaned_sections


def strip_duplicate_opening_heading_paragraph(cleaned_html: str, heading: str) -> str:
    heading = heading.strip()
    if not heading:
        return cleaned_html

    lines = [line for line in cleaned_html.splitlines() if line.strip()]
    if not lines:
        return cleaned_html

    match = re.fullmatch(r"<p>(.*?)</p>", lines[0].strip())
    if not match:
        return cleaned_html

    first_para = html.unescape(match.group(1)).strip()

    def normalize(value: str) -> str:
        return re.sub(r"\s+", " ", value).strip().casefold()

    if normalize(first_para) == normalize(heading):
        return "\n".join(lines[1:]).strip()

    return cleaned_html


def build_section_blocks(
    sections: List[Tuple[str, str]],
    *,
    drop_notes: bool,
    cleanup_settings: CleanupSettings | None = None,
) -> List[str]:
    blocks: List[str] = []
    first_real_section = True

    for heading, raw_html in sections:
        cleaned = sanitize_section_html(
            raw_html,
            drop_notes=drop_notes,
            cleanup_settings=cleanup_settings,
        )
        cleaned = strip_duplicate_opening_heading_paragraph(cleaned, heading)

        if not cleaned.strip():
            continue

        heading_html = (
            f'<h1 class="chapter-title">{html.escape(heading)}</h1>' if heading else ""
        )
        chapter_class = "chapter first-body-chapter" if first_real_section else "chapter"
        first_real_section = False

        blocks.append(
            f'<section class="{chapter_class}">{heading_html}{cleaned}</section>'
        )

    return blocks


def build_css(settings: LayoutSettings) -> str:
    mode = (settings.paragraph_spacing_mode or "traditional").strip().lower()

    if mode == "uniform":
        paragraph_css = """
p {
  margin: 0;
  text-align: justify;
  text-indent: 0;
  orphans: 2;
  widows: 2;
}

.chapter p + p {
  text-indent: 0;
  margin-top: 0;
}
"""
    elif mode == "no-indents":
        paragraph_css = """
p {
  margin: 0 0 0.65em 0;
  text-align: justify;
  text-indent: 0;
  orphans: 2;
  widows: 2;
}

.chapter p + p {
  text-indent: 0;
}
"""
    elif mode == "indented-compact":
        paragraph_css = """
p {
  margin: 0;
  text-align: justify;
  text-indent: 0;
  orphans: 2;
  widows: 2;
}

.chapter p + p {
  text-indent: 1.2em;
  margin-top: 0;
}
"""
    else:
        paragraph_css = """
p {
  margin: 0 0 0.65em 0;
  text-align: justify;
  text-indent: 0;
  orphans: 2;
  widows: 2;
}

.chapter p + p {
  text-indent: 1.2em;
}
"""

    return f"""
@page {{
  size: {cm(settings.trim_width_cm)} {cm(settings.trim_height_cm)};
  margin-top: {cm(settings.margin_top_cm)};
  margin-bottom: {cm(settings.margin_bottom_cm)};
  @bottom-center {{
    content: counter(page);
    font-size: 9pt;
  }}
}}

@page :right {{
  margin-left: {cm(settings.margin_inside_cm)};
  margin-right: {cm(settings.margin_outside_cm)};
}}

@page :left {{
  margin-left: {cm(settings.margin_outside_cm)};
  margin-right: {cm(settings.margin_inside_cm)};
}}

html {{
  font-size: {settings.font_size_pt}pt;
}}

body {{
  margin: 0;
  padding: 0;
  color: #111;
  font-family: {settings.font_family};
  line-height: {settings.line_height};
  -weasy-bookmark-level: none;
}}

.title-page {{
  break-after: page;
  min-height: 100%;
  display: flex;
  align-items: center;
  justify-content: center;
  text-align: center;
}}

.title-wrap h1 {{
  margin: 0 0 0.45em 0;
  font-size: {settings.font_size_pt * 2.0:.2f}pt;
  line-height: 1.15;
}}

.title-wrap .author {{
  margin: 0;
  font-size: {settings.font_size_pt * 1.15:.2f}pt;
}}

.blank-page {{
  break-after: page;
}}

.first-body-chapter {{
  break-before: right;
}}

.chapter {{
  break-before: page;
}}

.first-body-chapter.chapter {{
  break-before: right;
}}

.chapter-title {{
  text-align: center;
  margin: 0 0 1.8em 0;
  font-size: {settings.font_size_pt * 1.45:.2f}pt;
  page-break-after: avoid;
}}

{paragraph_css}

h1, h2, h3, h4 {{
  page-break-after: avoid;
}}

em {{
  font-style: italic;
}}

strong {{
  font-weight: bold;
}}

hr, .scene-break {{
  margin: 1.2em auto;
  width: 25%;
  border: 0;
  border-top: 1px solid #666;
}}
"""


def build_html_document(
    *,
    title: str,
    author: str,
    sections: List[Tuple[str, str]],
    settings: LayoutSettings,
    cleanup_settings: CleanupSettings | None = None,
) -> str:
    blocks = [
        f"""
<section class="title-page">
  <div class="title-wrap">
    <h1>{html.escape(title)}</h1>
    <p class="author">{html.escape(author)}</p>
  </div>
</section>
<div class="blank-page"></div>
""",
        *build_section_blocks(
            sections,
            drop_notes=settings.drop_notes,
            cleanup_settings=cleanup_settings,
        ),
    ]
    css = build_css(settings)

    return f"""
<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <title>{html.escape(title)}</title>
  <style>{css}</style>
</head>
<body>
{''.join(blocks)}
</body>
</html>
"""


def default_output_pdf_path(
    epub_path: Path,
    *,
    used_title: str,
    used_author: str,
) -> Path:
    book_slug = build_book_slug(
        title=used_title,
        author=used_author,
        fallback_stem=epub_path.stem,
    )
    output_dir = output_dir_for_book(epub_path.parent, book_slug)
    output_dir.mkdir(parents=True, exist_ok=True)
    return output_dir / interior_pdf_name(book_slug)


def default_output_docx_path(output_dir: Path, book_slug: str) -> Path:
    return output_dir / editable_docx_name(book_slug)


def default_output_markdown_path(output_dir: Path, book_slug: str) -> Path:
    return output_dir / markdown_name(book_slug)


def ensure_weasyprint_available() -> None:
    if HTML is None:
        raise RuntimeError(
            "WeasyPrint could not be imported because its native libraries are missing.\n"
            "Install WeasyPrint's system dependencies and try again.\n\n"
            f"Original import error: {WEASYPRINT_IMPORT_ERROR}"
        )


def render_pdf(html_doc: str, output_pdf_path: Path) -> None:
    ensure_weasyprint_available()
    output_pdf_path.parent.mkdir(parents=True, exist_ok=True)
    HTML(string=html_doc).write_pdf(str(output_pdf_path))

def _add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    text = OxmlElement("w:t")
    text.text = "1"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(text)
    run._r.append(fld_end)


def _configure_docx_section(section, settings: LayoutSettings) -> None:
    section.page_width = Cm(settings.trim_width_cm)
    section.page_height = Cm(settings.trim_height_cm)

    section.top_margin = Cm(settings.margin_top_cm)
    section.bottom_margin = Cm(settings.margin_bottom_cm)
    section.left_margin = Cm(settings.margin_inside_cm)
    section.right_margin = Cm(settings.margin_outside_cm)

    section.footer_distance = Cm(0.8)

    footer = section.footer
    footer_p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_number_field(footer_p)


def _apply_docx_paragraph_format(paragraph, settings: LayoutSettings, *, is_first_in_block: bool) -> None:
    fmt = paragraph.paragraph_format
    mode = (settings.paragraph_spacing_mode or "traditional").strip().lower()

    fmt.line_spacing = settings.line_height
    fmt.space_before = Pt(0)

    if mode == "uniform":
        fmt.space_after = Pt(0)
        fmt.first_line_indent = Pt(0)
    elif mode == "no-indents":
        fmt.space_after = Pt(settings.font_size_pt * 0.65)
        fmt.first_line_indent = Pt(0)
    elif mode == "indented-compact":
        fmt.space_after = Pt(0)
        fmt.first_line_indent = Pt(0 if is_first_in_block else settings.font_size_pt * 1.2)
    else:
        fmt.space_after = Pt(settings.font_size_pt * 0.65)
        fmt.first_line_indent = Pt(0 if is_first_in_block else settings.font_size_pt * 1.2)

    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def export_clean_docx(
    *,
    title: str,
    author: str,
    sections: List[Tuple[str, str]],
    output_path: str | Path | None,
    settings: LayoutSettings,
) -> Path | None:
    if output_path is None:
        return None

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Base document section
    section = doc.sections[0]
    _configure_docx_section(section, settings)

    # Normal style
    normal_style = doc.styles["Normal"]
    normal_style.font.size = Pt(settings.font_size_pt)

    # Title page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(settings.font_size_pt * 2.0)

    if author.strip():
        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author_para.add_run(author)
        author_run.font.size = Pt(settings.font_size_pt * 1.15)

    first_section = True

    for heading, text in sections:
        if first_section:
            doc.add_page_break()
            first_section = False
        else:
            new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
            _configure_docx_section(new_section, settings)

        if heading.strip():
            heading_para = doc.add_paragraph()
            heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            heading_run = heading_para.add_run(heading.strip())
            heading_run.bold = True
            heading_run.font.size = Pt(settings.font_size_pt * 1.45)
            heading_para.paragraph_format.space_after = Pt(settings.font_size_pt * 1.8)

        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
        first_body_paragraph = True

        for para in paragraphs:
            p = doc.add_paragraph()

            if is_scene_break(para) or para == "***":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run("***")
                run.bold = False
                p.paragraph_format.space_before = Pt(settings.font_size_pt * 1.2)
                p.paragraph_format.space_after = Pt(settings.font_size_pt * 1.2)
                p.paragraph_format.first_line_indent = Pt(0)
                first_body_paragraph = True
                continue

            run = p.add_run(para)
            run.font.size = Pt(settings.font_size_pt)
            _apply_docx_paragraph_format(
                p,
                settings,
                is_first_in_block=first_body_paragraph,
            )
            first_body_paragraph = False

    doc.save(str(output_path))
    return output_path


def _escape_markdown_text(text: str) -> str:
    return text.replace("\\", "\\\\").replace("*", "\\*").replace("_", "\\_")


def export_clean_markdown(
    *,
    title: str,
    author: str,
    sections: List[Tuple[str, str]],
    output_path: str | Path,
) -> Path:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    parts: list[str] = [f"# {_escape_markdown_text(title)}"]

    if author.strip():
        parts.append("")
        parts.append(f"_By {_escape_markdown_text(author)}_")

    for heading, text in sections:
        parts.append("")
        if heading.strip():
            parts.append(f"## {_escape_markdown_text(heading.strip())}")
            parts.append("")

        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
        for para in paragraphs:
            if is_scene_break(para) or para == "***":
                parts.append("***")
            else:
                parts.append(_escape_markdown_text(para))
            parts.append("")

    markdown_text = "\n".join(parts).rstrip() + "\n"
    output_path.write_text(markdown_text, encoding="utf-8")
    return output_path


def process_epub_to_pdf(
    *,
    epub_path: str | Path,
    output_pdf_path: str | Path | None = None,
    output_docx_path: str | Path | None = None,
    output_markdown_path: str | Path | None = None,
    export_docx: bool = False,
    export_markdown: bool = False,
    title: str | None = None,
    author: str | None = None,
    settings: LayoutSettings | None = None,
    cleanup_settings: CleanupSettings | None = None,
) -> EpubToPdfResult:
    epub_path = Path(epub_path)
    if not epub_path.exists():
        raise FileNotFoundError(f"EPUB not found: {epub_path}")

    settings = settings or LayoutSettings()
    cleanup_settings = cleanup_settings or CleanupSettings()
    epub_content = load_epub_content(epub_path)
    if not epub_content.sections:
        raise ValueError("No readable content sections found in EPUB.")

    used_title = title or epub_content.detected_title
    used_author = author or epub_content.detected_author

    html_doc = build_html_document(
        title=used_title,
        author=used_author,
        sections=epub_content.sections,
        settings=settings,
        cleanup_settings=cleanup_settings,
    )

    book_slug = build_book_slug(
        title=used_title,
        author=used_author,
        fallback_stem=epub_path.stem,
    )

    output_pdf = (
        Path(output_pdf_path)
        if output_pdf_path is not None
        else default_output_pdf_path(
            epub_path,
            used_title=used_title,
            used_author=used_author,
        )
    )

    render_pdf(html_doc, output_pdf)

    clean_text_sections: List[Tuple[str, str]] | None = None
    if export_docx or export_markdown:
        clean_text_sections = build_clean_text_sections(
            epub_content.sections,
            drop_notes=settings.drop_notes,
            cleanup_settings=cleanup_settings,
        )

    output_docx: Path | None = None
    if export_docx:
        output_docx = (
            Path(output_docx_path)
            if output_docx_path is not None
            else default_output_docx_path(output_pdf.parent, book_slug)
        )
        export_clean_docx(
            title=used_title,
            author=used_author,
            sections=clean_text_sections or [],
            output_path=output_docx,
            settings=settings,
        )

    output_markdown: Path | None = None
    if export_markdown:
        output_markdown = (
            Path(output_markdown_path)
            if output_markdown_path is not None
            else default_output_markdown_path(output_pdf.parent, book_slug)
        )
        export_clean_markdown(
            title=used_title,
            author=used_author,
            sections=clean_text_sections or [],
            output_path=output_markdown,
        )

    return EpubToPdfResult(
        input_epub=epub_path,
        output_dir=output_pdf.parent,
        output_pdf=output_pdf,
        output_docx=output_docx,
        output_markdown=output_markdown,
        book_slug=book_slug,
        detected_title=epub_content.detected_title,
        detected_author=epub_content.detected_author,
        used_title=used_title,
        used_author=used_author,
    )
