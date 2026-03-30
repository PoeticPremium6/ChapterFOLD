from __future__ import annotations

from pathlib import Path
from typing import List, Tuple

from docx import Document


def export_clean_docx(
    *,
    title: str,
    author: str,
    sections: List[Tuple[str, str]],
    output_path: str | Path,
) -> Path:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    doc.add_heading(title, level=0)
    if author.strip():
        doc.add_paragraph(author)

    first_section = True
    for heading, text in sections:
        if not first_section:
            doc.add_page_break()
        first_section = False

        if heading.strip():
            doc.add_heading(heading, level=1)

        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
        for para in paragraphs:
            doc.add_paragraph(para)

    doc.save(str(output_path))
    return output_path
